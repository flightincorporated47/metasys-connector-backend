from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.shared import Pt

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted, PageBreak
from reportlab.lib.units import inch

MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
MD_UL = re.compile(r"^\s*[-*]\s+(.*)$")
MD_OL = re.compile(r"^\s*(\d+)\.\s+(.*)$")


def read_text(p: Path) -> str:
    return p.read_text(encoding="utf-8", errors="replace")


def md_to_blocks(md: str) -> List[Tuple[str, str]]:
    blocks: List[Tuple[str, str]] = []
    lines = md.splitlines()
    in_code = False
    code_buf: List[str] = []
    para_buf: List[str] = []

    def flush_para():
        nonlocal para_buf
        text = "\n".join([x.strip() for x in para_buf]).strip()
        if text:
            blocks.append(("p", text))
        para_buf = []

    for line in lines:
        if line.strip().startswith("```"):
            if not in_code:
                flush_para()
                in_code = True
                code_buf = []
            else:
                blocks.append(("code", "\n".join(code_buf).rstrip()))
                in_code = False
            continue

        if in_code:
            code_buf.append(line.rstrip("\n"))
            continue

        m = MD_HEADING.match(line)
        if m:
            flush_para()
            level = len(m.group(1))
            blocks.append((f"h{level}", m.group(2).strip()))
            continue

        m = MD_UL.match(line)
        if m:
            flush_para()
            blocks.append(("ul", m.group(1).strip()))
            continue

        m = MD_OL.match(line)
        if m:
            flush_para()
            blocks.append(("ol", m.group(2).strip()))
            continue

        if not line.strip():
            flush_para()
            continue

        para_buf.append(line)

    flush_para()
    return blocks


def write_docx(md_path: Path, out_docx: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for kind, text in md_to_blocks(read_text(md_path)):
        if kind.startswith("h"):
            lvl = int(kind[1:])
            doc.add_heading(text, level=min(max(lvl, 1), 6))
        elif kind == "p":
            doc.add_paragraph(text)
        elif kind == "ul":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "ol":
            doc.add_paragraph(text, style="List Number")
        elif kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            doc.add_paragraph(text)

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)


def write_pdf(md_path: Path, out_pdf: Path) -> None:
    styles = getSampleStyleSheet()
    story = []

    for kind, text in md_to_blocks(read_text(md_path)):
        if kind.startswith("h"):
            lvl = int(kind[1:])
            style_name = {1: "Title", 2: "Heading1", 3: "Heading2", 4: "Heading3"}.get(lvl, "Heading3")
            story.append(Paragraph(text, styles[style_name]))
            story.append(Spacer(1, 0.12 * inch))
        elif kind == "p":
            story.append(Paragraph(text.replace("\n", "<br/>"), styles["BodyText"]))
            story.append(Spacer(1, 0.10 * inch))
        elif kind == "ul":
            story.append(Paragraph(f"• {text}", styles["BodyText"]))
        elif kind == "ol":
            story.append(Paragraph(text, styles["BodyText"]))
        elif kind == "code":
            story.append(Preformatted(text, styles["Code"]))
            story.append(Spacer(1, 0.12 * inch))

    out_pdf.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_pdf),
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    doc.build(story)


def write_combined_docx(md_files: List[Path], out_docx: Path) -> None:
    doc = Document()
    doc.add_heading("CSUDH Pilot — Connector SOPs", level=1)

    for i, p in enumerate(md_files):
        doc.add_heading(p.name, level=2)
        for kind, text in md_to_blocks(read_text(p)):
            if kind.startswith("h"):
                lvl = int(kind[1:])
                doc.add_heading(text, level=min(max(lvl + 1, 2), 6))
            elif kind == "p":
                doc.add_paragraph(text)
            elif kind == "ul":
                doc.add_paragraph(text, style="List Bullet")
            elif kind == "ol":
                doc.add_paragraph(text, style="List Number")
            elif kind == "code":
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.font.name = "Consolas"
                run.font.size = Pt(10)
        if i != len(md_files) - 1:
            doc.add_page_break()

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)


def write_combined_pdf(md_files: List[Path], out_pdf: Path) -> None:
    styles = getSampleStyleSheet()
    story = [Paragraph("CSUDH Pilot — Connector SOPs", styles["Title"]), Spacer(1, 0.2 * inch)]

    for i, p in enumerate(md_files):
        story.append(Paragraph(p.name, styles["Heading1"]))
        story.append(Spacer(1, 0.12 * inch))

        for kind, text in md_to_blocks(read_text(p)):
            if kind.startswith("h"):
                lvl = int(kind[1:])
                style_name = {1: "Heading1", 2: "Heading2", 3: "Heading3"}.get(lvl, "Heading3")
                story.append(Paragraph(text, styles[style_name]))
                story.append(Spacer(1, 0.10 * inch))
            elif kind == "p":
                story.append(Paragraph(text.replace("\n", "<br/>"), styles["BodyText"]))
                story.append(Spacer(1, 0.08 * inch))
            elif kind == "ul":
                story.append(Paragraph(f"• {text}", styles["BodyText"]))
            elif kind == "ol":
                story.append(Paragraph(text, styles["BodyText"]))
            elif kind == "code":
                story.append(Preformatted(text, styles["Code"]))
                story.append(Spacer(1, 0.10 * inch))

        if i != len(md_files) - 1:
            story.append(PageBreak())

    out_pdf.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_pdf),
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    doc.build(story)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--root", default=".", help="Repo root containing docs/SOPs")
    ap.add_argument("--out", default="exports", help="Output folder")
    ap.add_argument("--combine", action="store_true", help="Also produce combined files")
    args = ap.parse_args()

    root = Path(args.root).resolve()
    sop_root = root / "docs" / "SOPs"

    md_files = sorted([p for p in sop_root.rglob("*.md") if p.is_file()])
    if not md_files:
        raise SystemExit(f"No .md files found under: {sop_root}")

    out = Path(args.out).resolve()

    for md in md_files:
        rel = md.relative_to(sop_root)
        write_docx(md, out / "docx" / rel.with_suffix(".docx"))
        write_pdf(md, out / "pdf" / rel.with_suffix(".pdf"))

    if args.combine:
        write_combined_docx(md_files, out / "combined" / "CSUDH_Connector_SOPs.docx")
        write_combined_pdf(md_files, out / "combined" / "CSUDH_Connector_SOPs.pdf")

    print(f"Exported {len(md_files)} SOP markdown files into: {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
